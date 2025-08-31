from __future__ import annotations
from pathlib import Path
from typing import Iterable
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from core.logger import logger


# ---------- helpers ----------
def _ensure_font(paragraph, name="Segoe UI", size=11, bold=False, color: tuple[int,int,int]=(0,0,0)):
    """Apply font to all runs inside a paragraph."""
    if not paragraph.runs:
        run = paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(*color)

def _add_spacer(doc: Document, height_pt: int = 4):
    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(height_pt)
    return p

def _add_hr(doc: Document):
    """Add a thin horizontal rule using a single-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    tc = t.rows[0].cells[0]
    tc.width = Inches(6.0)
    # set bottom border
    tc._tc.get_or_add_tcPr().append(_borders_xml(bottom=(1, "000000")))
    # remove paragraph content height
    tc.paragraphs[0].clear()

def _borders_xml(top=None, left=None, bottom=None, right=None):
    """
    Build w:tcBorders or w:tblBorders part.
    Each side is (size_pt, hexcolor) -> size in 1/8 pt per Word spec (2 = 0.25pt).
    """
    borders = OxmlElement('w:tcBorders')
    for side_name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        if val:
            sz_pt, color = val
            side = OxmlElement(f"w:{side_name}")
            side.set(qn("w:val"), "single")
            side.set(qn("w:sz"), str(int(max(1, sz_pt*8))))   # Word expects 1/8 pt units
            side.set(qn("w:color"), color)
            borders.append(side)
    return borders

def _tbl_borders_xml(sz=1, color="1E2230"):
    borders = OxmlElement('w:tblBorders')
    for side in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(int(max(1, sz*8))))
        e.set(qn("w:color"), color)
        borders.append(e)
    return borders

def _bullet_block(doc: Document, title: str, items: Iterable[str]):
    doc.add_heading(title, level=2)
    for it in (list(items) or ["Chưa xác định"]):
        p = doc.add_paragraph(it, style=None)
        p.style = doc.styles["List Bullet"] if "List Bullet" in doc.styles else None
        p.paragraph_format.space_after = Pt(2)
        _ensure_font(p, size=11)

def _apply_page_margins(doc: Document, left=0.8, right=0.8, top=0.7, bottom=0.7):
    """inches"""
    for section in doc.sections:
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)


# ---------- main API ----------
def export_meeting_report(docx_path: str, meta: dict, summary: str, extracted: dict):
    """
    Export a polished meeting report.
    `extracted` keys: goal, agenda, attendance, decisions, action_items (list of dict)
    """
    logger.info("export_meeting_report -> %s", docx_path)
    Path(docx_path).parent.mkdir(parents=True, exist_ok=True)

    # Document + base styles
    doc = Document()
    _apply_page_margins(doc)
    # Base font
    if "Normal" in doc.styles:
        doc.styles["Normal"].font.name = "Segoe UI"
        doc.styles["Normal"].font.size = Pt(11)

    # Title
    title_text = meta.get("title", "Meeting Report")
    title_p = doc.add_paragraph(title_text)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _ensure_font(title_p, size=22, bold=True)

    # Subtitle meta (datetime)
    dt = meta.get("datetime") or ""
    sub_p = doc.add_paragraph(dt)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _ensure_font(sub_p, size=10, color=(100,114,139))  # muted gray
    _add_spacer(doc, 2)

    # Summary (paragraph block)
    doc.add_heading("Summary", level=2)
    p = doc.add_paragraph(summary or "Chưa xác định")
    p.paragraph_format.space_after = Pt(4)
    _ensure_font(p, size=11)
    _add_spacer(doc, 6)

    # Goal / Agenda / Attendance / Decisions (bulleted)
    _bullet_block(doc, "Goal", extracted.get("goal", []))
    _add_spacer(doc, 4)
    _bullet_block(doc, "Agenda", extracted.get("agenda", []))
    _add_spacer(doc, 4)
    _bullet_block(doc, "Meeting Attendance", extracted.get("attendance", []))
    _add_spacer(doc, 4)
    _bullet_block(doc, "Decisions", extracted.get("decisions", []))
    _add_spacer(doc, 6)

    # Action Items — proper table
    doc.add_heading("Action Items", level=2)
    ai = extracted.get("action_items") or []
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # table borders
    tbl_pr = table._tbl.tblPr
    tbl_pr.append(_tbl_borders_xml(sz=1, color="1E2230"))

    hdr = table.rows[0].cells
    hdr[0].text = "Action Item"
    hdr[1].text = "Assigned To"
    hdr[2].text = "Due Date"
    hdr[3].text = "Completed"

    # header styling
    for i, c in enumerate(hdr):
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < 3 else WD_ALIGN_PARAGRAPH.CENTER
            _ensure_font(p, size=11, bold=True)
        # cell borders (thicker bottom for header)
        c._tc.get_or_add_tcPr().append(_borders_xml(bottom=(1, "1E2230")))

    if not ai:
        row = table.add_row().cells
        row[0].text = "Chưa xác định"
        row[1].text = ""
        row[2].text = ""
        row[3].text = ""
    else:
        for it in ai:
            row = table.add_row().cells
            row[0].text = it.get("item", "") or ""
            row[1].text = it.get("assignee", "") or ""
            row[2].text = it.get("due", "") or ""
            row[3].text = "✔" if it.get("done") else ""

    # body rows styling
    for r_idx, row in enumerate(table.rows[1:], start=1):
        for c_idx, c in enumerate(row.cells):
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx < 3 else WD_ALIGN_PARAGRAPH.CENTER
                _ensure_font(p, size=11)
            # thin borders for body
            c._tc.get_or_add_tcPr().append(_borders_xml(top=(1, "1E2230"), bottom=(1, "1E2230")))

    # Optional footer or note
    _add_spacer(doc, 10)
    foot = doc.add_paragraph("Generated by AI Meeting Assistant")
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _ensure_font(foot, size=9, color=(100,114,139))

    # Save
    doc.save(docx_path)
    logger.info("export_meeting_report saved -> %s", docx_path)
